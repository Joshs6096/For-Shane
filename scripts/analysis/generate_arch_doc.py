from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Set narrow margins
from docx.shared import Inches
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# Styles
styles = doc.styles

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text='', bold=False, italic=False):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_table_from_rows(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    doc.add_paragraph()

# ── Title ──────────────────────────────────────────────────────────────
title = doc.add_heading('BC-Remittance / Cur8 AI — Infrastructure & Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Compiled by BC Engineering Assistant   |   Date: 2026-05-10').italic = True
doc.add_paragraph()

# ── 1. Overview ────────────────────────────────────────────────────────
add_heading('1. Overview', 1)
add_para('Cur8 AI maintains two related codebases:')
add_table_from_rows(
    ['Repo', 'Purpose', 'Stage'],
    [
        ['BC-Remittance\n(GitHub: Joshs6096/BC-Remittance)',
         'Full-stack remittance automation platform — first customer is Black Clover',
         'Active production build'],
        ['cur8-platform\n(local: /Users/josh/code/cur8-platform)',
         'Shared multi-tenant API base platform that BC-Remittance derives from',
         'Foundation / shared scaffold'],
    ]
)
add_para('Both use the same runtime stack (FastAPI · SQLAlchemy async · PostgreSQL) by design. '
         'Future Cur8 customers will use cur8-platform as the starting point; BC-specific logic lives in BC-Remittance.')

# ── 2. BC-Remittance Architecture ─────────────────────────────────────
add_heading('2. BC-Remittance Architecture', 1)

add_heading('2.1 Services', 2)
add_table_from_rows(
    ['Service', 'Technology', 'Port', 'Role'],
    [
        ['API (bc-remittance-api)', 'Python 3.11 · FastAPI · uvicorn', '8000', 'REST API, JWT auth, ingest upload, admin'],
        ['Worker (bc-remittance-worker)', 'Python 3.11 · asyncio', '—', 'Async pipeline: extraction → matching → feedback consumers'],
        ['Web (bc-remittance-web)', 'React 18 · TypeScript · Vite · TanStack Query · Tailwind', '5173', 'AR reviewer portal'],
        ['Postgres (bc-remittance-postgres)', 'PostgreSQL 16', '5432', 'Primary database — all relational state'],
        ['Redis (bc-remittance-redis)', 'Redis 7', '6379', 'Redis Streams message bus (worker queue)'],
        ['Nginx (prod only)', 'nginx', '443 / 80', 'TLS termination + reverse proxy → API + Web'],
    ]
)

add_heading('2.2 Backend Layers', 2)
add_code_block(
    "backend/app/\n"
    "├── api/\n"
    "│   ├── routes/remittance.py      # PDF/XLSX/CSV/EDI upload endpoint (ingest)\n"
    "│   └── v1/                       # v1 REST API routes\n"
    "│       ├── aliases.py            # Customer alias CRUD\n"
    "│       ├── deposits.py           # Click-to-deposit queue\n"
    "│       ├── documents.py          # Document management\n"
    "│       ├── health.py             # /health probe\n"
    "│       ├── matches.py            # Match review (approve/reject)\n"
    "│       ├── remittances.py        # Remittance CRUD\n"
    "│       ├── reports.py            # AR reporting\n"
    "│       └── tenants.py            # Tenant config (ADMIN role only)\n"
    "├── services/\n"
    "│   ├── ocr/                      # OCR providers (hybrid: text + Claude Haiku)\n"
    "│   ├── normalization/            # Customer alias resolution, invoice cleanup\n"
    "│   ├── matching/\n"
    "│   │   ├── tier1_exact.py        # Exact invoice match\n"
    "│   │   ├── tier2_shortpay.py     # Variance / short-pay match\n"
    "│   │   ├── tier3_subset_sum.py   # Aggregated subset-sum match\n"
    "│   │   └── tier4_heuristic.py    # Heuristic / ML match\n"
    "│   ├── execution/                # NetSuite write-back\n"
    "│   │   ├── customer_payment.py\n"
    "│   │   ├── deposit.py\n"
    "│   │   └── credit_memo.py\n"
    "│   └── netsuite/                 # NS SuiteTalk REST client (TBA auth, pool, write-back)\n"
    "├── workers/\n"
    "│   ├── extraction_consumer.py\n"
    "│   ├── matching_consumer.py\n"
    "│   └── feedback_consumer.py\n"
    "├── models/                       # SQLAlchemy ORM models\n"
    "└── core/\n"
    "    ├── config.py                 # Pydantic settings (env vars)\n"
    "    └── logging.py                # structlog setup"
)

add_heading('2.3 Redis Stream Pipeline', 2)
add_table_from_rows(
    ['Topic', 'Published by', 'Consumed by', 'Meaning'],
    [
        ['remittance.ingested', 'Ingestion service (API)', 'ExtractionConsumer', 'New doc uploaded, OCR pending'],
        ['remittance.normalized', 'ExtractionConsumer', 'MatchingConsumer', 'Lines extracted + aliases resolved'],
        ['remittance.matched', 'MatchingConsumer', 'API / UI', 'Match result ready for review'],
        ['remittance.execution', 'API (match approved)', 'ExecutionConsumer', 'Approved match → NS write'],
        ['remittance.deposited', 'ExecutionConsumer', '—', 'NS Customer Payment posted'],
        ['match.approved', 'API (reviewer action)', 'FeedbackConsumer', 'Human approved match'],
        ['match.rejected', 'API (reviewer action)', 'FeedbackConsumer', 'Human rejected match'],
    ]
)
add_para('Dead-letter queues (.dlq) hold messages that exhaust 5 retries.')

add_heading('2.4 External Integrations', 2)
add_table_from_rows(
    ['Integration', 'Protocol', 'Auth', 'Concurrency limit'],
    [
        ['NetSuite SuiteTalk REST', 'HTTPS REST', 'TBA — HMAC-SHA256 OAuth 1.0 M2M', '5 simultaneous connections per tenant (semaphore in pool.py)'],
        ['Anthropic Claude Haiku', 'HTTPS REST', 'API key (ANTHROPIC_API_KEY)', 'OCR fallback when PDF text is sparse (<100 chars)'],
    ]
)

add_heading('2.5 Auth Model', 2)
add_table_from_rows(
    ['Role', 'Level', 'Permissions'],
    [
        ['READ_ONLY', '1', 'Read remittances, matches, deposits, reports'],
        ['REVIEWER', '2', '+ Upload docs, approve/reject matches, create deposits'],
        ['AR_MANAGER', '3', '+ Admin reporting'],
        ['ADMIN', '4', '+ Tenant config management'],
    ]
)
add_para('JWT (symmetric, jwt_secret env var). All endpoints require Bearer token except /api/v1/health.')

# ── 3. cur8-platform ───────────────────────────────────────────────────
add_heading('3. cur8-platform Architecture', 1)
add_para('Minimal multi-tenant API scaffold used as the shared foundation for future Cur8 AI customers.')
add_table_from_rows(
    ['Service', 'Technology', 'Port', 'Role'],
    [
        ['API', 'Python 3.11 · FastAPI', '8000', 'REST API (tenant CRUD + hello smoke-test)'],
        ['Postgres', 'PostgreSQL 16', '5432', 'Primary DB'],
    ]
)
add_para('No Redis, no frontend, no workers. Purely API + DB. Migrations via Alembic.')

# ── 4. Connection Diagram ─────────────────────────────────────────────
add_heading('4. Connection & Dependency Diagram', 1)
add_para('(Text representation — see Mermaid source in Paperclip for interactive rendering)')
add_code_block(
    "External actors\n"
    "  Browser (AR Reviewer) ──HTTPS──► Nginx (prod, 443)\n"
    "                                       ├──proxy──► Web (React, :5173)\n"
    "                                       └──proxy /api──► API (FastAPI, :8000)\n"
    "  Remittance PDFs / XLSX / CSV / EDI ──upload──► API\n"
    "\n"
    "BC-Remittance internal data flows\n"
    "  API ──read/write──► PostgreSQL 16 (:5432)\n"
    "  API ──XADD remittance.ingested──► Redis 7 (:6379)\n"
    "  Worker ──XREAD consumer group──► Redis\n"
    "  Worker ──read/write──► PostgreSQL\n"
    "  Worker ──XADD result topics──► Redis\n"
    "\n"
    "OCR path\n"
    "  Worker ──PDF bytes (text-first)──► Anthropic Claude Haiku\n"
    "\n"
    "NetSuite execution\n"
    "  Worker ──SuiteTalk REST TBA──► NetSuite\n"
    "  API ──SuiteTalk REST TBA──► NetSuite\n"
    "\n"
    "cur8-platform\n"
    "  cur8 API ──read/write──► cur8 PostgreSQL (:5432)\n"
    "  BC-Remittance API ··shares FastAPI/SQLAlchemy runtime·· cur8 API"
)

# ── 5. End-to-End Data Flow ────────────────────────────────────────────
add_heading('5. End-to-End Data Flow', 1)
add_code_block(
    "AR uploads PDF\n"
    "      │\n"
    "      ▼\n"
    "POST /api/v1/remittances  (API)\n"
    "      │  stores doc, publishes remittance.ingested\n"
    "      ▼\n"
    "ExtractionConsumer  (Worker)\n"
    "      │  text extract → Claude Haiku OCR fallback\n"
    "      │  normalise lines, resolve customer aliases\n"
    "      │  publishes remittance.normalized\n"
    "      ▼\n"
    "MatchingConsumer  (Worker)\n"
    "      │  Tier 1: exact invoice match\n"
    "      │  Tier 2: variance / short-pay\n"
    "      │  Tier 3: subset-sum aggregated\n"
    "      │  Tier 4: heuristic ML\n"
    "      │  high-confidence → publishes remittance.execution\n"
    "      │  low-confidence  → publishes remittance.matched (→ review queue)\n"
    "      ▼\n"
    "AR reviews in portal (if needed)\n"
    "      │  POST /api/v1/matches/{id}/approve or /reject\n"
    "      │  publishes match.approved / match.rejected\n"
    "      ▼\n"
    "FeedbackConsumer  (Worker)  ←── alias-learning loop\n"
    "      │\n"
    "ExecutionConsumer / API  (writes to NetSuite)\n"
    "      │  POST Customer Payment\n"
    "      │  AR clicks Deposit → POST Deposit\n"
    "      │  Write-off lines → Credit Memo\n"
    "      ▼\n"
    "NetSuite (source of truth — GL entries posted)"
)

# ── 6. Environments ────────────────────────────────────────────────────
add_heading('6. Infrastructure — Environments', 1)
add_table_from_rows(
    ['Env', 'Compose file', 'Notes'],
    [
        ['Local / dev', 'docker-compose.yml', 'Hot-reload, local volumes, port-mapped services'],
        ['Production', 'docker-compose.prod.yml', 'Nginx TLS, env from .env.prod, resource limits, restart: unless-stopped. Target OS: Windows Server (Docker Desktop)'],
    ]
)
add_para('Production resource limits: Postgres 1 GB RAM / 1 CPU · Redis 256 MB (maxmemory allkeys-lru)')

# ── 7. Key Environment Variables ───────────────────────────────────────
add_heading('7. Key Environment Variables', 1)
add_table_from_rows(
    ['Variable', 'Service', 'Purpose'],
    [
        ['DATABASE_URL', 'API, Worker', 'asyncpg connection string'],
        ['REDIS_URL', 'API, Worker', 'Redis connection'],
        ['ANTHROPIC_API_KEY', 'Worker', 'Claude Haiku OCR'],
        ['ANTHROPIC_MODEL', 'Worker', 'Default claude-haiku-4-5-20251001'],
        ['OCR_PROVIDER', 'API, Worker', 'dual (text + Claude), textract (legacy), etc.'],
        ['NS_ACCOUNT_ID', 'API, Worker', 'NetSuite account (per-tenant, also stored in DB)'],
        ['NS_CLIENT_ID / NS_PRIVATE_KEY_PATH', 'Worker', 'TBA auth for NetSuite'],
        ['JWT_SECRET', 'API', 'JWT signing key'],
        ['STORAGE_BACKEND', 'API', 'local (dev) or s3 (prod)'],
        ['WEB_CONCURRENCY', 'API', 'uvicorn worker count (prod: 2×CPU+1)'],
    ]
)

# ── 8. Relationship ────────────────────────────────────────────────────
add_heading('8. cur8-platform vs BC-Remittance Relationship', 1)
add_code_block(
    "cur8-platform\n"
    "├─ Tenant model + multi-tenancy pattern\n"
    "├─ FastAPI / SQLAlchemy async base\n"
    "├─ Docker + Alembic migration pattern\n"
    "└─ Shared runtime config conventions\n"
    "\n"
    "BC-Remittance  (extends / implements cur8 spec)\n"
    "├─ All of the above +\n"
    "├─ OCR pipeline (text + Claude Haiku hybrid)\n"
    "├─ Redis Streams 5-layer queue\n"
    "├─ NetSuite SuiteTalk integration\n"
    "├─ 4-tier match engine\n"
    "├─ React AR reviewer portal\n"
    "└─ Nginx production infra"
)
add_para('BC-Remittance implements the Cur8 spec (§2 layers, §3.3/§4.5/§5.6 topics, §7.1 NS write rules). '
         'Future Cur8 customers extend cur8-platform the same way BC-Remittance does.')

output_path = '/tmp/BC_Remittance_Architecture.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
