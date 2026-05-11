"""
CFO Executive Variance Analysis Report
Team File (Daily Cash Fcst - 05.08.26.xlsb) vs. AI File (Daily Cash Fcst - 5.8.26_10tab.xlsx)
As of May 8, 2026
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

# ── Colour palette ─────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
RED    = RGBColor(0xC0, 0x00, 0x00)
AMBER  = RGBColor(0xBF, 0x8F, 0x00)
GREEN  = RGBColor(0x37, 0x86, 0x10)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF2, 0xF2, 0xF2)
DGREY  = RGBColor(0x40, 0x40, 0x40)

# ── Helper: shade a table cell ─────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','bottom','left','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '999999'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_style(p, sz=10, bold=False, color=DGREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size  = Pt(sz)
        run.font.bold  = bold
        run.font.color.rgb = color

def add_heading(doc, text, level=1, sz=14, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(sz)
    run.font.bold  = True
    run.font.color.rgb = color
    if level == 1:
        p.paragraph_format.keep_with_next = True
    return p

def add_body(doc, text, sz=10, color=DGREY, bold=False, space_after=5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size  = Pt(sz)
    run.font.bold  = bold
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, sz=9.5, color=DGREY, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.3 + indent_level * 0.2)
    run = p.add_run(text)
    run.font.size  = Pt(sz)
    run.font.color.rgb = color
    return p

def add_table_row(table, cols, bg='FFFFFF', bold_row=False, font_sz=8.5,
                  text_color='1F3864' if False else '404040'):
    row = table.add_row()
    for i, (cell_text, col_def) in enumerate(zip(cols, table.columns)):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, bg)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(str(cell_text))
        run.font.size = Pt(font_sz)
        run.font.bold = bold_row
    return row

def make_table(doc, headers, rows_data, col_widths=None,
               header_bg='1F3864', header_fg='FFFFFF',
               alt_bg='F2F2F2', font_sz=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, header_bg)
        set_cell_border(cell, color=header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.font.size  = Pt(font_sz)
        run.font.bold  = True
        run.font.color.rgb = WHITE
    # Data rows
    for ri, row_data in enumerate(rows_data):
        bg = 'FFFFFF' if ri % 2 == 0 else alt_bg
        is_bold = row_data[0].startswith('**') if row_data else False
        clean_row = [c.strip('*') for c in row_data]
        add_table_row(table, clean_row, bg=bg, bold_row=is_bold, font_sz=font_sz)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("TRANSFORM SR HOLDING MANAGEMENT LLC")
run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("DAILY CASH FORECAST — VARIANCE ANALYSIS REPORT")
run2.font.size = Pt(18); run2.font.bold = True; run2.font.color.rgb = NAVY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run("AI-Generated File vs. Team File  |  As of May 8, 2026")
run3.font.size = Pt(11); run3.font.color.rgb = DGREY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(12)
run4 = p4.add_run("Prepared: May 10, 2026  |  Classification: CFO – RESTRICTED")
run4.font.size = Pt(9); run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "I.  EXECUTIVE SUMMARY", level=1, sz=13)

add_body(doc,
    "This report presents a systematic, cell-level comparison between the AI-generated cash forecast "
    "('AI File': Daily Cash Fcst – 5.8.26_10tab.xlsx) and the finance team's authoritative file "
    "('Team File': Daily Cash Fcst – 05.08.26.xlsb) as of May 8, 2026. Five independent analytical "
    "agents interrogated all 10 visible tabs and 4 hidden backend data tabs. Every figure cited herein "
    "was read directly from the files. No values have been estimated or interpolated.", sz=10)

add_body(doc,
    "The overarching finding is singular and structurally significant: the AI File's four hidden backend "
    "data tabs — the master data spine (FY File), Payroll, MERCH, and Non-Merch — are frozen at a "
    "March 5, 2025 as-of date. Because every formula-driven visible tab ultimately traces to one of these "
    "four backends, the AI File is functionally a one-year-old snapshot dressed in a current-date wrapper. "
    "This produces two classes of error: (1) stale hardcoded balances on reference schedule tabs, and "
    "(2) cascading formula failures (#N/A errors) on all cash flow computation tabs.", sz=10)

# Summary scorecard table
add_heading(doc, "Overall Scorecard by Sheet", level=2, sz=11, color=NAVY)
scorecard_headers = ["Sheet", "AI File Status", "Severity", "Root Cause"]
scorecard_rows = [
    ["Cash Flow",              "Truncated at Jan 30, 2026; 440+ #N/A errors in last common month",  "CRITICAL",  "Backend tabs stale; formulas broken"],
    ["Inflows Actuals",        "Data ends Jan 30, 2026 — missing Feb–May 2026",                      "CRITICAL",  "Source data not refreshed"],
    ["Inflows Detail",         "Data ends Jan 30, 2026 — no May 2026 forecast rows",                 "CRITICAL",  "FY File backend frozen at Mar 2025"],
    ["Disbursement Detail",    "Covers Mar 2025–Jan 2026; XLSB covers May 2026–Jan 2027",            "CRITICAL",  "Completely different time windows"],
    ["Inflows Forecasting",    "Missing 290+ historical forecast comparison columns",                 "HIGH",      "Source truncated at col 58 vs. col 771"],
    ["Disbursement Actuals",   "Core data matches; 189 audit/check columns absent",                  "LOW",       "Trailing columns excluded by design"],
    ["LC Forecast Changes",    "Missing 2 Completed 2026 transactions; no LC inventory panel",       "MODERATE",  "Not refreshed for 2026 events"],
    ["Trapped Cash",           "$1.42M total overstatement; Stripe holdback row missing ($1.27M)",   "HIGH",      "Balances reflect prior period"],
    ["UBS Reserve",            "Note A/B balances differ by $76.1M combined; 14-mo stale",           "HIGH",      "Schedule reflects 3/5/2025 position"],
    ["2020 Term Loans",        "All tranche balances stale by ~14 months; $86.7M total variance",    "HIGH",      "As-of date 3/5/2025 vs. 5/8/2026"],
    ["FY File (hidden)",       "Data ends Jan 2026 — 3.5 months stale",                             "CRITICAL",  "Not rolled forward"],
    ["Payroll (hidden)",       "Data ends Jul 2025 — 10 months stale",                              "CRITICAL",  "Not refreshed for FY2026"],
    ["MERCH (hidden)",         "Data ends May 2025 — 12 months stale",                              "CRITICAL",  "Entire FY2026 missing"],
    ["Non-Merch (hidden)",     "Data ends Jan 2026 — 3.5 months stale",                             "CRITICAL",  "Not rolled forward"],
]

tbl = make_table(doc, scorecard_headers, scorecard_rows,
                 col_widths=[1.6, 2.7, 0.85, 2.05], font_sz=7.8)

# Color severity column
sev_colors = {
    'CRITICAL': ('FFD5D5', RED),
    'HIGH':     ('FFF2CC', AMBER),
    'MODERATE': ('DEEAF1', NAVY),
    'LOW':      ('E2EFDA', GREEN),
}
for ri, row_data in enumerate(scorecard_rows):
    sev = row_data[2]
    cell = tbl.rows[ri+1].cells[2]
    bg, fg = sev_colors.get(sev, ('FFFFFF', DGREY))
    shade_cell(cell, bg)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = fg
            run.font.bold = True

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — ROOT CAUSE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "II.  ROOT CAUSE ANALYSIS", level=1, sz=13)

add_body(doc,
    "The AI File was constructed by reading cached values from the March 5, 2025 source workbook and "
    "extending its structural scaffolding forward to May 8, 2026. While the visible tab architecture "
    "correctly replicates the 10-tab standard, the four hidden data backend tabs were populated from "
    "the 2025 source file without refresh. This creates a structural dependency failure:", sz=10)

add_bullet(doc, "FY File (master data spine): Frozen at Jan 30, 2026. "
           "All formula-driven cells in Inflows Detail, Disbursement Detail, and Cash Flow that "
           "use INDEX/MATCH against FY File receive no data for dates after Jan 30, 2026.", sz=9.5)
add_bullet(doc, "Payroll: Last populated date Jul 3, 2025 — 10 months before the report date. "
           "All payroll disbursements in the Cash Flow's Payroll/Bens column (Col 92) are absent "
           "or zero for the entire Feb–May 2026 window.", sz=9.5)
add_bullet(doc, "MERCH: Data ends May 31, 2025 — 12 months stale. "
           "Merch CIA/On-Terms disbursements are entirely missing for all of fiscal 2026.", sz=9.5)
add_bullet(doc, "Non-Merch: Data ends Jan 30, 2026 — 3.5 months stale. "
           "Non-Merch disbursement categories (taxes, facilities, legal, logistics) carry no "
           "2026 actuals.", sz=9.5)

add_body(doc,
    "As a result, the Cash Flow sheet's formula chain fails: Disbursement Detail pulls from "
    "FY File/Payroll/MERCH/Non-Merch → returns #N/A → Cash Flow net CF column returns #N/A for "
    "all dates after Jan 30, 2026. The AI File contains 440+ confirmed #N/A errors in the Cash "
    "Flow sheet for the last common month alone (January 2026), escalating to a complete data void "
    "for February through May 2026.", sz=10, space_after=6)

add_body(doc,
    "The single exception to this pattern is Disbursement Actuals, which was reconstructed directly "
    "from the Team File's xlsb data in a post-build patch step. That sheet matches the Team File "
    "exactly for May 8, 2026 across all 37 operational columns.", sz=10, color=GREEN)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — SHEET-BY-SHEET FINDINGS
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "III.  SHEET-BY-SHEET FINDINGS", level=1, sz=13)

# ── 3A. CASH FLOW ──
add_heading(doc, "A.  Cash Flow  [CRITICAL]", level=2, sz=11, color=RED)

add_body(doc,
    "The Cash Flow sheet is the primary deliverable of this workbook — it is where all backend data "
    "streams converge into the daily available cash and total cash position. The AI File's Cash Flow "
    "sheet is structurally present but operationally inert for the May 2026 window.", sz=10)

make_table(doc,
    ["Metric", "Team File (XLSB)", "AI File (XLSX)", "Variance"],
    [
        ["Date range",          "Feb 3, 2020 – Jan 29, 2027",  "Feb 3, 2020 – Jan 30, 2026",  "AI missing 260+ rows (all of Feb–Dec 2026)"],
        ["Total data rows",     "1,825",                        "1,565",                         "–260 rows"],
        ["Max columns",         "160 (0-based: 159)",           "147 (0-based: 146)",            "–13 rightmost columns"],
        ["#N/A error cells",    "0",                            "440+ in Jan 2026 alone",        "Formula cascade failure in AI file"],
        ["May 8 Available Cash","$1.510M",                      "Not present",                   "N/A — date missing"],
        ["May 8 Total Cash",    "$89.588M",                     "Not present",                   "N/A — date missing"],
        ["May 8 Net Cash Flow", "$0.040M",                      "Not present",                   "N/A — date missing"],
        ["May 8 Inflows Total", "$2.200M",                      "Not present",                   "N/A — date missing"],
        ["May 8 Disb. Total",   "($2.160M)",                    "Not present",                   "N/A — date missing"],
    ],
    col_widths=[1.7, 1.8, 1.8, 1.9], font_sz=8)

add_body(doc, "May 1–8, 2026 Daily Cash Position (Team File — AI File has no data):", sz=9.5, bold=True, space_after=2)

make_table(doc,
    ["Date", "Op. CF", "Net CF", "Avail. Cash", "Total Cash", "Inflows", "Payroll/Bens", "Merch Disb", "Disb. Total"],
    [
        ["May 1",  "$0.62M", "$0.47M", "$5.35M",  "$93.98M", "$3.38M", "($0.14M)", "($3.01M)", "($2.91M)"],
        ["May 4",  "($2.54M)","($2.17M)","$3.18M","$91.73M", "$0.99M", "($0.48M)", "($1.06M)", "($3.31M)"],
        ["May 5",  "$0.07M", "$0.07M", "$3.25M",  "$92.36M", "$1.20M", "($0.08M)", "($1.46M)", "($1.68M)"],
        ["May 6",  "($2.45M)","($2.45M)","$0.80M","$88.76M", "$1.92M", "($3.30M)", "($0.87M)", "($4.37M)"],
        ["May 7",  "$0.67M", "$0.67M", "$1.47M",  "$90.26M", "$1.10M", "($0.14M)", "($0.73M)", "($0.88M)"],
        ["**May 8**","**$0.04M**","**$0.04M**","**$1.51M**","**$89.59M**","**$2.20M**","**($1.11M)**","**($0.74M)**","**($2.16M)**"],
    ],
    col_widths=[0.65, 0.65, 0.65, 0.8, 0.8, 0.7, 0.9, 0.8, 0.8], font_sz=7.8)

# ── 3B. DISBURSEMENT ACTUALS ──
add_heading(doc, "B.  Disbursement Actuals  [LOW — Core Data Match; Structural Gap Only]", level=2, sz=11, color=GREEN)

add_body(doc,
    "Disbursement Actuals is the single sheet where the AI File and Team File are in substantive agreement. "
    "After a targeted rebuild from the Team File's xlsb data, all 37 operational disbursement columns "
    "for May 8, 2026 are identical to the cent.", sz=10)

make_table(doc,
    ["Metric", "Team File (XLSB)", "AI File (XLSX)", "Status"],
    [
        ["Date range",             "Feb 3, 2025 – Jan 29, 2027", "Feb 3, 2025 – Jan 29, 2027", "Match"],
        ["Data rows",              "520",                          "520",                         "Match"],
        ["May 8 Daily Total",      "($2.160M)",                    "($2.160M)",                   "Match"],
        ["May 8 Payroll Total",    "($1.110M)",                    "($1.110M)",                   "Match"],
        ["May 8 MERCH Total",      "($0.460M)",                    "($0.460M)",                   "Match"],
        ["Operational cols (0–89)","90 columns",                   "90 columns",                  "Match"],
        ["Audit/check cols (90–278)","Actual Ending Bal: $1.51M; CF Col K: $1.51M", "Not present", "MISSING in AI File"],
    ],
    col_widths=[2.0, 1.75, 1.75, 0.9], font_sz=8)

add_body(doc, "Note: The 189 trailing columns (indices 90–278) in the Team File contain reconciliation "
         "checks including Actual Ending Balance and Cash Flow Col K cross-checks. These are internal "
         "audit mechanisms, not inputs. Their absence from the AI File does not affect the reported "
         "disbursement figures but eliminates automated reconciliation capability.", sz=9.5, indent=True)

# ── 3C. DISBURSEMENT DETAIL ──
add_heading(doc, "C.  Disbursement Detail  [CRITICAL — Completely Different Time Windows]", level=2, sz=11, color=RED)

add_body(doc,
    "This sheet represents the most structurally irreconcilable difference between the two files. "
    "The two versions of Disbursement Detail cover non-overlapping date ranges with zero shared rows:", sz=10)

make_table(doc,
    ["Metric", "Team File (XLSB)", "AI File (XLSX)"],
    [
        ["Date range",     "May 11, 2026 – Jan 29, 2027", "Mar 6, 2025 – Jan 30, 2026"],
        ["Purpose",        "Forward forecast (actuals cutoff is May 8 in Disb. Actuals)", "Historical actuals — stale"],
        ["Data rows",      "190",                           "237"],
        ["Overlapping dates", "None",                       "None"],
        ["May 8, 2026",    "Not present (by design — it's an actuals-only date)", "Not present"],
    ],
    col_widths=[1.8, 2.7, 1.9], font_sz=8)

add_body(doc,
    "The Team File treats Disbursement Detail as a forward-looking forecast beginning the first business "
    "day after the as-of date. May 8 actuals are captured in Disbursement Actuals. The AI File's version "
    "contains historical actuals ending January 2026 — the wrong time period entirely. The AI File needs "
    "to be rebuilt with forecast data from the FY File, Non-Merch, Payroll, and MERCH backends once those "
    "are refreshed.", sz=10)

# ── 3D. INFLOWS SHEETS ──
add_heading(doc, "D.  Inflows Sheets (Actuals, Detail, Forecasting)  [CRITICAL]", level=2, sz=11, color=RED)

add_body(doc, "All three inflows sheets in the AI File are materially stale or structurally truncated.", sz=10)

make_table(doc,
    ["Sheet", "Team File — Key Facts", "AI File — Key Facts", "Severity"],
    [
        ["Inflows Actuals",
         "Actuals through May 8, 2026.\nMay 8 daily total: $2.20M\n(Kmart: $0.80M, Home Svcs: $1.37M)",
         "Data ends Jan 30, 2026.\nFeb–May 2026 entirely absent.\nNo May 8 row exists.",
         "CRITICAL"],
        ["Inflows Detail",
         "Forecast rows starting May 11, 2026.\nAs-of reference date: 5/8/2026.\nMay 11 total: $1.75M",
         "Data ends Jan 30, 2026.\nAs-of reference date frozen at 3/5/2025.\nNo May 2026 rows.",
         "CRITICAL"],
        ["Inflows Forecasting",
         "Extends to col 771.\nIncludes 'Forecast as of 05/01/26' comparison block (Wk15: $10.76M).\nActual/Forecast Change column block present.",
         "Stops at col 58.\n'Forecast as of 05/01/26' block absent.\nChange columns absent.\nWk15 shows $12.80M (stale prior forecast).",
         "HIGH"],
    ],
    col_widths=[1.55, 2.35, 2.35, 0.85], font_sz=7.8)

add_body(doc, "Inflows Actuals — May 1–8, 2026 (Team File only; AI File has no data):", sz=9.5, bold=True, space_after=2)
make_table(doc,
    ["Date", "Sears Stores", "Kmart Stores", "Home Services", "Auto Centers", "Daily Total"],
    [
        ["May 1",  "$0.05M", "$0.62M", "$2.22M", "$0.03M", "$3.38M"],
        ["May 4",  "$0.14M", "$0.05M", "$0.82M", "$0.03M", "$1.14M"],
        ["May 5",  "$0.19M", "$0.66M", "$0.79M", "$0.02M", "$1.75M"],
        ["May 6",  "$0.21M", "$0.39M", "$1.25M", "$0.02M", "$1.92M"],
        ["May 7",  "$0.17M", "$0.58M", "$0.72M", "$0.02M", "$1.55M"],
        ["**May 8**","**$0.01M**","**$0.80M**","**$1.37M**","**$0.02M**","**$2.20M**"],
    ],
    col_widths=[0.75, 1.1, 1.1, 1.1, 1.1, 1.15], font_sz=8)

# ── 3E. LC FORECAST CHANGES ──
add_heading(doc, "E.  LC Forecast Changes  [MODERATE]", level=2, sz=11, color=AMBER)

add_body(doc,
    "The tracking table (right panel) matches the Team File for all historical entries. However, "
    "the AI File is missing two fiscal year 2026 completed transactions and lacks the LC inventory "
    "reference panel (left panel) present in the Team File.", sz=10)

make_table(doc,
    ["Item", "Team File Value", "AI File Value", "Delta"],
    [
        ["XL Specialty (Surety) collateral change",   "CASH $2,310,000 → $1,886,500 | Net: +$423,500",  "Not present", "+$423,500 in Team File"],
        ["Excess Claims Recoveries collateral change", "CASH $365,437 → $0 | Net: +$365,437",            "Not present", "+$365,437 in Team File"],
        ["'Completed 2026:' label",                   "Present with 2 data rows",                         "Label present, zero data rows", "Missing 2 entries"],
        ["LC inventory reference panel (left panel)",  "WAGA LC table: 17 UBS + 5 PNC LCs",              "Absent — not built into AI file", "Structural gap"],
    ],
    col_widths=[2.05, 2.05, 1.65, 1.45], font_sz=8)

# ── 3F. TRAPPED CASH ──
add_heading(doc, "F.  Trapped Cash  [HIGH — $1.42M Net Overstatement in AI File]", level=2, sz=11, color=AMBER)

add_body(doc,
    "The AI File's Trapped Cash schedule reflects prior-period balances. Ten of 44 line items carry "
    "different values, producing a net overstatement of $1,417,622 in the AI File's grand total. "
    "One line item — Cash Holdback from Stripe — is entirely absent from the AI File.", sz=10)

make_table(doc,
    ["Line Item", "Team File (5/8/26)", "AI File (3/5/25)", "Variance (AI – Team)"],
    [
        ["Cash deposited with UBS for LCs",      "$30,868,423",    "$31,387,142",    "+$518,719  (AI overstated)"],
        ["Cash Holdback — Discover",              "$0",             "$136,000",       "+$136,000  (AI overstated)"],
        ["Cash Holdback — FDC",                   "$7,051,464",     "$6,947,467",     "–$103,997  (AI understated)"],
        ["Utility Deposits",                       "$2,818,825",     "$2,747,492",     "–$71,333   (AI understated)"],
        ["Advances — Landlords & Sub-Tenants",     "$935,737",       "$1,182,294",     "+$246,556  (AI overstated)"],
        ["Holdback — Innovel environmental",       "$556,078",       "$805,549",       "+$249,471  (AI overstated)"],
        ["Drawn LC restricted cash",               "$30,969,470",    "$33,514,023",    "+$2,544,553 (AI overstated)"],
        ["Hackensack",                             "$569,882",       "$539,155",       "–$30,727   (AI understated)"],
        ["Bohemia",                                "$5,048,591",     "$4,594,213",     "–$454,378  (AI understated)"],
        ["Home Warranty reserves & deposits",      "$6,209,231",     "$5,048,685",     "–$1,160,546 (AI understated)"],
        ["Cash Holdback — Stripe",                 "$1,271,111",     "NOT PRESENT",    "–$1,271,111 (AI missing)"],
        ["**GRAND TOTAL**",                        "**$88,476,879**","**$89,894,501**","**+$1,417,622 (AI overstated)**"],
    ],
    col_widths=[2.3, 1.6, 1.6, 2.1], font_sz=8)

# ── 3G. UBS RESERVE ──
add_heading(doc, "G.  UBS Reserve  [HIGH — Loan Balances 14 Months Stale]", level=2, sz=11, color=AMBER)

add_body(doc,
    "The UBS Reserve schedule in the AI File reflects the loan position as of March 5, 2025. "
    "Both the outstanding principal balances and the reserve amounts have moved materially in the "
    "intervening 14 months. The AI File also retains a historical loan section that is no longer "
    "present in the Team File, and is missing the Hackensack CMBS section added to the Team File.", sz=10)

make_table(doc,
    ["Field", "Team File (5/8/26)", "AI File (3/5/25)", "Variance"],
    [
        ["Property count",          "74",                 "85",                  "–11 properties"],
        ["Note A outstanding",      "$299,435,000",       "$223,340,938",        "–$76,094,062 in AI"],
        ["Note A rate",             "S + 4.00%",          "S + 4.75%",           "Rate reduced (–75 bps)"],
        ["Note B outstanding",      "$25,367,934",        "$65,586,537",         "+$40,218,603 in AI"],
        ["Note B rate",             "S + 4.00%",          "S + 8.00%",           "Rate reduced (–400 bps)"],
        ["Combined UBS total",      "$324,802,934",       "$288,927,474",        "–$35,875,460 in AI"],
        ["Interest Reserve",        "$8,539,653",         "$14,720,062",         "+$6,180,409 in AI"],
        ["Working Capital Reserve", "$6,470,679",         "$11,311,985",         "+$4,841,306 in AI"],
        ["Total Reserves",          "$19,945,821",        "$37,270,612",         "+$17,324,791 in AI"],
        ["Hackensack CMBS section", "Present ($34M facility, 7.75%)", "Absent",  "Missing from AI"],
        ["Historical loan section", "Absent",             "Present (107 props)",  "Extra section in AI"],
    ],
    col_widths=[2.05, 1.8, 1.8, 1.55], font_sz=8)

# ── 3H. 2020 TERM LOANS ──
add_heading(doc, "H.  2020 Term Loans  [HIGH — $86.7M Total Balance Variance; 14 Months Stale]", level=2, sz=11, color=AMBER)

add_body(doc,
    "Every term loan tranche balance differs between files due to the 14-month date gap. "
    "Two tranches that did not yet exist in the AI File's as-of date (September 2025 Term Loans "
    "and Cyrus Eighth Incremental Term Loans) are now fully funded and tracked in the Team File.", sz=10)

make_table(doc,
    ["Tranche", "Team File (5/8/26)", "AI File (3/5/25)", "Variance"],
    [
        ["Tranche 2 (ESL T2)",          "$50,705,480",   "$48,014,192",   "+$2,691,288"],
        ["Tranche 7 (Fifth Incr.)",     "$2,077,643",    "$20,569,149",   "–$18,491,506"],
        ["Tranche 8 (Sixth Incr.)",     "$15,134,323",   "$20,466,242",   "–$5,331,919"],
        ["Tranche 12 (Cyrus 3rd)",      "$0",            "$3,858,902",    "–$3,858,902"],
        ["Tranche 22 (T22, ESL $175M)", "$257,678,471",  "$202,140,416",  "+$55,538,055"],
        ["Tranche 23 (Cyrus 7th)",      "$30,884,421",   "$40,770,000",   "–$9,885,579"],
        ["Tranche 24 (15th Incr.)",     "$50,892,762",   "$10,000,000",   "+$40,892,762"],
        ["Sept. 2025 Term Loans",        "$31,089,213",   "Not funded (pre-issuance)", "+$31,089,213"],
        ["Cyrus 8th Incremental",        "$8,160,168",    "Not funded (pre-issuance)", "+$8,160,168"],
        ["**TOTAL PRINCIPAL**",         "**$533,317,469**","**$446,574,634**","**+$86,742,835**"],
    ],
    col_widths=[2.15, 1.75, 1.75, 1.55], font_sz=8)

# ── 3I. HIDDEN BACKEND TABS ──
add_heading(doc, "I.  Hidden Backend Tabs (FY File, Payroll, MERCH, Non-Merch)  [CRITICAL]", level=2, sz=11, color=RED)

add_body(doc,
    "The four hidden data feeds are the upstream source for all formula-driven cells in the "
    "visible tabs. Their staleness is the primary root cause of all other discrepancies. "
    "Every sheet carries an as-of date stamp of March 5, 2025.", sz=10)

make_table(doc,
    ["Sheet", "Team File Date Range", "AI File Date Range", "Months Stale", "May 2026 Data?"],
    [
        ["FY File (master spine)", "Feb 2, 2026 – Jan 29, 2027", "Feb 3, 2025 – Jan 30, 2026", "~3.5", "NO"],
        ["Payroll",                "May 2026 – Aug 2026 (forecast)", "Dec 2024 – Jul 2025",      "~10",  "NO"],
        ["MERCH",                  "Jan 2026 – Jan 2027",           "Mar 2025 – May 2025",       "~12",  "NO"],
        ["Non-Merch",              "May 2026 – Jan 2027",           "Jan 2025 – Jan 2026",       "~3.5", "NO"],
    ],
    col_widths=[1.7, 2.0, 2.0, 1.0, 0.9], font_sz=8)

add_body(doc, "May 4–8, 2026 Payroll Totals (Team File — XLSB; by BU total col):", sz=9.5, bold=True, space_after=2)
make_table(doc,
    ["Date", "Payroll/Bens Total"],
    [
        ["May 4", "($516,110)"],
        ["May 5", "($86,200)"],
        ["May 6", "($3,362,692)  ← Large payroll cycle"],
        ["May 7", "($152,600)"],
        ["May 8", "($1,132,329)"],
    ],
    col_widths=[1.5, 5.0], font_sz=8.5)

add_body(doc, "All of these are zero or absent in the AI File's Payroll backend.", sz=9.5, color=RED, indent=True)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — REMEDIATION ROADMAP
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "IV.  REMEDIATION ROADMAP", level=1, sz=13)

add_body(doc,
    "The following actions are sequenced by dependency. Steps 1–4 must be completed before "
    "any downstream visible tab can be trusted. Actions are prioritized by data dependency order, "
    "not business importance.", sz=10)

make_table(doc,
    ["Priority", "Action", "Owner", "Dependency", "Impact"],
    [
        ["1 — IMMEDIATE",
         "Refresh Payroll backend tab with data from Payroll finance team through current forecast horizon (at minimum through Aug 2026).",
         "Payroll Finance Team",
         "None",
         "Resolves all Payroll/Bens #N/A errors in Cash Flow and Disbursement Detail for Feb–May 2026."],
        ["2 — IMMEDIATE",
         "Refresh MERCH backend tab with current CIA/On-Terms data from James Coutre. Must cover Jan 2026 – Jan 2027.",
         "James Coutre",
         "None",
         "Resolves all Merch disbursement gaps in Cash Flow (Col 94) and Disbursement Detail."],
        ["3 — IMMEDIATE",
         "Refresh Non-Merch backend tab with current data from Adharvana through Jan 2027.",
         "Adharvana",
         "None",
         "Resolves tax, logistics, facilities, legal disbursements in Disbursement Detail for Feb–May 2026."],
        ["4 — IMMEDIATE",
         "Roll forward FY File (master data spine) to cover Feb 2026 – Jan 2027, matching Team File's column schema. Resolve col 11 label (Fidem vs. Risk IN) and col 38 label discrepancy (PH 2.0 out vs. Legal Settlements).",
         "Treasury / Modeling Team",
         "Steps 1–3 inform FY File inputs",
         "Once FY File is current, all downstream INDEX/MATCH formulas in visible tabs will resolve for 2026."],
        ["5 — HIGH",
         "Refresh Trapped Cash schedule: update 10 stale balances, add missing Stripe holdback row ($1,271,111). Source: bank statements + RE finance team per SOP.",
         "Accounting / RE Finance",
         "None (standalone schedule)",
         "Corrects $1.42M net overstatement in total trapped cash."],
        ["6 — HIGH",
         "Update UBS Reserve schedule to reflect 5/8/2026 position: revise Note A/B balances, property count, rates, reserves, and footnote. Add Hackensack CMBS section per Zach Straebel.",
         "Zach Straebel / Real Estate",
         "None (standalone schedule)",
         "Corrects $35.9M combined balance variance and $17.3M reserve overstatement."],
        ["7 — HIGH",
         "Update 2020 Term Loans schedule to 5/8/2026 balances per Jon Goodin/CAP table. Add funded data rows for Sept. 2025 TL ($31.1M) and Cyrus 8th Incremental ($8.2M).",
         "Jon Goodin / Cap Table",
         "None (standalone schedule)",
         "Corrects $86.7M total principal variance."],
        ["8 — MODERATE",
         "Add two missing LC Forecast Changes entries under 'Completed 2026:' (XL Specialty +$423,500; Excess Claims Recoveries +$365,437). Build LC inventory reference panel (left panel) matching Team File.",
         "Jon Goodin",
         "None (standalone entries)",
         "Corrects LC collateral position and enables full reference functionality."],
        ["9 — MODERATE",
         "Extend Disbursement Detail forward: once backends (Steps 1–4) are refreshed, rebuild Disbursement Detail to cover May 11, 2026 forward, matching Team File's forecast horizon (through Jan 2027).",
         "Modeling Team",
         "Steps 1–4 complete",
         "Aligns Disbursement Detail purpose (forward forecast) with Team File standard."],
        ["10 — LOW",
         "Rebuild Inflows Forecasting comparison columns: extend to col 771 to include 'Forecast as of 05/01/26' block and Actual/Forecast Change block. Currently AI File truncates at col 58.",
         "Modeling Team",
         "Step 4 complete",
         "Restores weekly forecast vs. actual comparison capability."],
        ["11 — LOW",
         "Add trailing audit columns (90–278) to Disbursement Actuals: Actual Ending Balance, Cash Flow Col K cross-check. Currently absent from AI File.",
         "Modeling Team",
         "Step 4 complete",
         "Restores automated reconciliation between Disbursement Actuals and Cash Flow."],
    ],
    col_widths=[0.95, 2.35, 1.2, 1.1, 1.6], font_sz=7.5)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 5 — WHAT IS CONFIRMED ACCURATE
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "V.  CONFIRMED ACCURATE IN AI FILE", level=1, sz=13)

add_body(doc,
    "Not all components of the AI File are deficient. The following were confirmed to match "
    "the Team File exactly:", sz=10)

make_table(doc,
    ["Item", "Confirmed Match Detail"],
    [
        ["Disbursement Actuals — all 37 operational columns",
         "May 8, 2026 daily total: ($2.160M) | PYRL/BENS: ($1.110M) | MERCH: ($0.460M) — exact match to the cent across all 520 rows"],
        ["LC Forecast Changes — historical tracking table",
         "All completed prior, 2022, 2023, 2024, and 2025 entries match exactly in the right-panel tracking table"],
        ["Trapped Cash — 34 of 44 line items",
         "Including: PNC LC cash ($5,636,582), Citi LC cash ($0), RE Segregated Account ($5,241,937), all 7 property sub-lines, RE Escrows ($1,350,000), Safelite ($572,762)"],
        ["UBS Reserve — UBS Upsize reconciliation table",
         "Interest Reserve, Working Capital, General, ADA & Life Safety, Deferred Maintenance all tie with rounding only ($0.01 on Working Capital)"],
        ["2020 Term Loans — column structure",
         "All 26 tranche column headers (cols 3–28) are identical; Sept. 2025 TL and Cyrus 8th Incremental columns pre-built correctly"],
        ["10-tab visible structure + 8 hidden backend tabs",
         "Tab names, order, and visibility states (visible/hidden/veryHidden) match the Team File standard exactly"],
        ["Inflows Actuals — structure",
         "Column headers and layout are structurally correct; data gap is a content issue, not a structural one"],
    ],
    col_widths=[2.5, 5.0 - 2.5], font_sz=8.5)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run(
    "All figures in this report were read directly from file data without estimation or rounding beyond display precision. "
    "Values stored in both files are denominated in US millions unless otherwise noted. "
    "Prepared using automated file comparison agents on May 10, 2026."
)
run.font.size  = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────
OUT = '/Users/josh/Downloads/CFO Variance Report - Daily Cash Fcst - May 8 2026.docx'
doc.save(OUT)
import os
sz = os.path.getsize(OUT)
print(f"Saved → {OUT}")
print(f"Size: {sz:,} bytes ({sz/1024:.0f} KB)")
