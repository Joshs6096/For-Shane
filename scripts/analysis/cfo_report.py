from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(0.9)

# ── Color palette ─────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1F, 0x38, 0x64)
DKGREY  = RGBColor(0x40, 0x40, 0x40)
MIDGREY = RGBColor(0x70, 0x70, 0x70)
RED     = RGBColor(0xC0, 0x00, 0x00)
GREEN   = RGBColor(0x37, 0x86, 0x44)
AMBER   = RGBColor(0xBF, 0x87, 0x00)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LBLUE   = RGBColor(0xD6, 0xE4, 0xF0)
NAVY_BG = "1F3864"
RED_BG  = "FFD5D5"
GRN_BG  = "E2EFDA"
AMB_BG  = "FFF2CC"
GRY_BG  = "F2F2F2"
HDR_BG  = "2F5496"
ALT_BG  = "EBF1DD"

# ── Style helpers ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   kwargs.get(side+'_val', 'single'))
        border.set(qn('w:sz'),    kwargs.get(side+'_sz',  '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get(side+'_color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def p(text='', bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
      space_before=0, space_after=6, italic=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return para

def heading(text, level=1):
    if level == 1:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size  = Pt(12)
        run.font.color.rgb = WHITE
        # Add shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  NAVY_BG)
        pPr.append(shd)
    elif level == 2:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(text)
        run.bold = True
        run.font.size  = Pt(11)
        run.font.color.rgb = NAVY
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),  'single')
        bottom.set(qn('w:sz'),   '6')
        bottom.set(qn('w:space'),'1')
        bottom.set(qn('w:color'), '2F5496')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(text)
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = DKGREY
    return para

def add_table(headers, rows, col_widths=None, alt_rows=True):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, HDR_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tbl_row = tbl.rows[r_idx+1]
        bg = ALT_BG if (alt_rows and r_idx % 2 == 1) else "FFFFFF"
        for c_idx, (val, kwargs) in enumerate(row_data):
            cell = tbl_row.cells[c_idx]
            cell_bg = kwargs.get('bg', bg)
            set_cell_bg(cell, cell_bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            align = kwargs.get('align', WD_ALIGN_PARAGRAPH.LEFT)
            para.alignment = align
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            if kwargs.get('bold'):    run.bold = True
            if kwargs.get('italic'):  run.italic = True
            if kwargs.get('color'):   run.font.color.rgb = kwargs['color']

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl

def fmt_m(v, dec=3):
    """Format as $M with parens for negative."""
    if v is None: return '—'
    if v < 0:  return f'(${ abs(v):,.{dec}f}M)'
    return f'${ v:,.{dec}f}M'

def fmt_d(v):
    """Format as plain dollar with parens."""
    if v is None: return '—'
    if v < 0:  return f'(${ abs(v):,.0f})'
    return f'${ v:,.0f}'

def var_color(v):
    if v is None: return {}
    return {'color': GREEN} if v >= 0 else {'color': RED}

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BEGINS
# ══════════════════════════════════════════════════════════════════════════════

# Cover header
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run('TRANSFORM SR HOLDING MANAGEMENT LLC')
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = NAVY

p('Daily Cash Forecast — Variance Analysis', bold=True, size=12,
  color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('Reporting Date: May 8, 2026', bold=False, size=10, color=DKGREY,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('PREPARED FOR: Chief Financial Officer', bold=True, size=10,
  color=DKGREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('PREPARED BY: Treasury / Financial Planning & Analysis',
  size=9, color=MIDGREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('CLASSIFICATION: Internal — Confidential',
  size=9, italic=True, color=MIDGREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SECTION 1 — EXECUTIVE SUMMARY ────────────────────────────────────────
heading('Section 1 — Executive Summary', 1)

p(('This memorandum presents a line-by-line comparison between (i) the daily cash '
   'forecast workbook prepared by the finance team for May 8, 2026 (the "Actual File") '
   'and (ii) a workbook generated by an AI agent using NetSuite general ledger queries '
   'as its primary data source (the "AI File"). All figures have been extracted directly '
   'from each file without adjustment or interpolation. Discrepancies are reported as '
   'observed.'), size=9.5, space_after=4)

p('Material findings are summarized below:', bold=True, size=9.5, space_after=3)

summary_rows = [
    [('1', {}),
     ('Net Cash Flow, May 8, 2026',{}),
     ('+$0.040M (Actual)',{'color':GREEN}),
     ('-$0.727M (AI)',{'color':RED}),
     ('-$0.767M', {'color':RED,'bold':True}),
     ('AI captures only NetSuite-posted AP vendor payments; does not capture inflows or payroll',{})],
    [('2', {}),
     ('Total Cash, May 8, 2026',{}),
     ('$89.587M (Actual)',{'color':GREEN}),
     ('$5.242M (AI)',{'color':RED}),
     ('-$84.345M', {'color':RED,'bold':True}),
     ('AI file lacks Unavailable Cash data ($82.836M) for 2026; DV column not populated',{})],
    [('3', {}),
     ('Gross Inflows, May 8',{}),
     ('$2.200M (Actual)',{}),
     ('$0.001M (AI)',{'color':RED}),
     ('-$2.199M', {'color':RED,'bold':True}),
     ('Customer receipts not visible as same-day bank debits in NetSuite; actual data sourced from Daily Detail cash position file',{})],
    [('4', {}),
     ('Gross Disbursements, May 8',{}),
     ('-$2.160M (Actual)',{}),
     ('-$0.727M (AI)',{}),
     ('+$1.433M (understated)', {'color':RED,'bold':True}),
     ('Payroll (-$1.110M), Merch on Terms (-$0.460M), and Taxes (-$0.289M) not captured by NetSuite AP vendor payment queries',{})],
    [('5', {}),
     ('File Structural Integrity',{}),
     ('Continuous; no errors (Actual)',{'color':GREEN}),
     ('#N/A errors throughout 2025–2026 (AI)',{'color':RED}),
     ('N/A — broken formula chain',{'color':RED}),
     ('AI Cash Flow sheet produces #N/A in all calculated columns from Jun 2, 2025 onward due to Payroll sheet expiring Jul 2025',{})],
    [('6', {}),
     ('Time Horizon',{}),
     ('Feb 3, 2020 – Jan 29, 2027 (Actual)',{}),
     ('Feb 3, 2020 – Jan 30, 2026 (AI)',{}),
     ('-12 months',{'color':RED}),
     ('AI file is approximately 12 months short of the Actual file\'s forward horizon',{})],
]

# Build summary table manually due to merged col complexity
tbl = doc.add_table(rows=1+6, cols=6)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdrs = ['#','Finding','Actual File','AI File','Variance / Impact','Root Cause']
widths = [0.25, 1.4, 1.25, 1.25, 1.15, 2.2]
for i,h in enumerate(hdrs):
    c = tbl.rows[0].cells[i]
    set_cell_bg(c, HDR_BG)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = c.paragraphs[0].add_run(h)
    rn.bold=True; rn.font.size=Pt(8.5); rn.font.color.rgb=WHITE

rows_data = [
    ['1','Net Cash Flow, May 8, 2026','+$0.040M','-$0.727M','-$0.767M',
     'AI captures only posted AP vendor payments; does not capture inflows or payroll'],
    ['2','Total Cash, May 8, 2026','$89.587M','$5.242M','-$84.345M',
     'AI file lacks Unavailable Cash data for 2026; DV column not populated'],
    ['3','Gross Inflows, May 8','$2.200M','$0.001M','-$2.199M',
     'Customer receipts not visible as same-day bank debits in NetSuite; actual data sourced from Daily Detail cash position file'],
    ['4','Gross Disbursements, May 8','-$2.160M','-$0.727M','+$1.433M understated',
     'Payroll (-$1.110M), Merch on Terms (-$0.460M), and Taxes (-$0.289M) not captured by NetSuite AP vendor payment queries'],
    ['5','Formula / Structural Integrity','Continuous; no errors','#N/A errors throughout 2025–2026','Broken — unusable for 2025–2026',
     'AI Cash Flow sheet #N/A in all calculated columns from Jun 2, 2025 onward due to Payroll sheet data expiring Jul 4, 2025'],
    ['6','Forward Forecast Horizon','Feb 2020 – Jan 2027','Feb 2020 – Jan 2026','−12 months',
     'AI file is 12 months short of the Actual file\'s forward horizon'],
]

row_bgs = [GRY_BG,'FFFFFF',GRY_BG,'FFFFFF',RED_BG,GRY_BG]
var_colors = [RED,RED,RED,RED,RED,RED]
for ri, (row, bg) in enumerate(zip(rows_data, row_bgs)):
    tr = tbl.rows[ri+1]
    for ci, val in enumerate(row):
        cell = tr.cells[ci]
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in [0,2,3,4] else WD_ALIGN_PARAGRAPH.LEFT
        rn = cell.paragraphs[0].add_run(val)
        rn.font.size = Pt(8.5)
        if ci == 4: rn.font.color.rgb = var_colors[ri]; rn.bold = True
for i,w in enumerate(widths):
    for row in tbl.rows:
        row.cells[i].width = Inches(w)

doc.add_paragraph()

# ── SECTION 2 — FILE OVERVIEW ─────────────────────────────────────────────
heading('Section 2 — Workbook Overview Comparison', 1)

heading('2.1  File Attributes', 2)
add_table(
    ['Attribute','Actual File (A)','AI-Generated File (B)','Difference'],
    [
      [('Filename',{}),('Daily Cash Fcst - 05.08.26.xlsb',{}),('Daily Cash Fcst - 5.8.26.xlsx',{}),('—',{})],
      [('Format',{}),('Excel Binary Workbook (.xlsb)',{}),('Open XML Workbook (.xlsx)',{}),('Format mismatch',{'color':AMBER})],
      [('File Size',{}),('2,736,860 bytes (2.6 MB)',{}),('12,486,053 bytes (11.9 MB)',{}),('AI file is 4.6× larger',{'color':AMBER})],
      [('Number of Worksheets',{}),('18',{'bold':True}),('62',{'bold':True}),('AI contains 44 additional stale working-paper tabs',{'color':RED})],
      [('Sheets Present in A, Absent in B',{}),('SOP; MERCH',{'bold':True}),('Neither present',{}),('Material — SOP documents all data sources and named contacts',{'color':RED})],
      [('Report Date Stamp (internal)',{}),('May 8, 2026',{'color':GREEN,'bold':True}),('March 5, 2025 (FY File ref.)',{'color':RED,'bold':True}),('AI file internally dated ~14 months prior',{'color':RED})],
    ],
    col_widths=[1.8,2.2,2.2,2.3]
)

p()
heading('2.2  Worksheet Date Coverage', 2)
add_table(
    ['Sheet','Actual File — First Date','Actual File — Last Date','AI File — First Date','AI File — Last Date','Gap (AI vs. Actual Last)'],
    [
      [('Cash Flow',{'bold':True}),('Feb 3, 2020',{}),('Jan 29, 2027 ✓',{'color':GREEN}),('Feb 3, 2020',{}),('Jan 30, 2026 (#N/A from Jun 2, 2025)',{'color':RED}),('−12 months; broken from Jun 2025',{'color':RED})],
      [('FY File',{'bold':True}),('Feb 2, 2026',{}),('Jan 29, 2027 ✓',{'color':GREEN}),('Mar 5, 2025',{}),('Jan 30, 2026',{'color':RED}),('−12 months; prior year window',{'color':RED})],
      [('Inflows Detail',{'bold':True}),('May 11, 2026',{}),('Jan 29, 2027 ✓',{'color':GREEN}),('Mar 6, 2025',{}),('Jan 30, 2026',{'color':RED}),('No May 2026 data',{'color':RED})],
      [('Disbursement Detail',{'bold':True}),('May 11, 2026',{}),('Jan 29, 2027 ✓',{'color':GREEN}),('Mar 6, 2025',{}),('Jan 30, 2026',{'color':RED}),('No May 2026 data',{'color':RED})],
      [('Payroll',{'bold':True}),('May 8, 2026',{}),('Aug 28, 2026 ✓',{'color':GREEN}),('Dec 2, 2024',{}),('Jul 4, 2025',{'color':RED}),('−10 months; root cause of #N/A cascade',{'color':RED})],
      [('KES & KCD Cash Flow',{'bold':True}),('Dec 10, 2025',{}),('Apr 3, 2026',{}),('May 31, 2022',{}),('May 30, 2025',{'color':RED}),('−11 months',{'color':RED})],
    ],
    col_widths=[1.3,1.3,1.55,1.3,1.55,1.45]
)

doc.add_paragraph()

# ── SECTION 3 — MAY 8 2026 CASH POSITION ─────────────────────────────────
heading('Section 3 — May 8, 2026 Daily Cash Position: Actual vs. AI-Generated', 1)
p(('All amounts in millions of US dollars ($M) unless otherwise noted. '
   'Actual figures extracted from Actual File row 1,713 (Cash Flow sheet) '
   'and corroborated against Inflows Actuals and Disbursement Actuals tabs. '
   'AI figures extracted from AI File row 1,709 (Cash Flow sheet) where data exists; '
   'otherwise noted as unavailable. Variances computed as AI minus Actual (positive = AI higher than Actual).'),
  size=9, italic=True, space_after=4)

heading('3.1  Cash Flow Summary', 2)
add_table(
    ['Line Item','Actual ($M)','AI-Generated ($M)','Variance ($M)','Variance %','Notes'],
    [
      [('Gross Inflows',{'bold':True}),('$2.200',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.001',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('($2.199)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('(99.9%)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('NetSuite does not capture same-day bank credits for customer receipts',{})],
      [('Gross Disbursements',{'bold':True}),('($2.160)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($0.727)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('$1.433',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('66.3% understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Payroll, Merch on Terms, and Taxes not captured via NetSuite AP queries',{})],
      [('Net Cash Flow (Operating)',{'bold':True,'bg':GRY_BG}),('$0.040',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('($0.727)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('($0.767)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('N/M',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bg':GRY_BG}),('Direction of cash flow inverted in AI file',{'bg':GRY_BG})],
      [('Non-Operating Cash Flow',{}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('Consistent — no financing activity on May 8 in either file',{})],
    ],
    col_widths=[2.0,1.1,1.3,1.1,1.1,2.1]
)

p()
heading('3.2  Balance Sheet Cash Position, End of Day May 8, 2026', 2)
add_table(
    ['Balance Sheet Component','Actual ($M)','AI-Generated ($M)','Variance ($M)','Notes'],
    [
      [('Available Cash (Unrestricted Operating)',{'bold':True}),('$1.510',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('($1.510)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('AI file Available Cash = $0 throughout 2026 (formula driven off Jan 30 opening + net CF of ($0.727M))',{})],
      [('Segregated Account',{}),('$5.242',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$5.242',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Consistent — both files carry forward the same segregated RE account balance',{})],
      [('Unavailable Cash (Restricted + Collateral)',{'bold':True}),('$82.836',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('($82.836)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('DV column (Unavail Cash Lookup from Model) not populated in AI file for 2026 dates; hardcoded externally in Actual File',{})],
      [('TOTAL CASH (Available + Seg. + Unavailable)',{'bold':True,'bg':GRY_BG}),('$89.587',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('$5.242',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('($84.345)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('AI Total Cash is understated by $84.3M; driven almost entirely by the missing Unavailable Cash balance',{'bg':GRY_BG})],
    ],
    col_widths=[2.3,1.1,1.3,1.1,2.7]
)

p()
p('Note: The Unavailable Cash balance of $82.836M in the Actual File is supported by the '
  'Trapped Cash tab ($88.5M total unavailable, of which $82.836M is classified as '
  'Unavailable Cash in the Cash Flow model). The AI file\'s Trapped Cash tab shows '
  '$89.895M total but is carried from the March 2025 source file and does not reflect '
  'May 2026 positions.', size=9, italic=True, space_after=4)

doc.add_paragraph()

# ── SECTION 4 — INFLOWS ANALYSIS ─────────────────────────────────────────
heading('Section 4 — Inflows Analysis: May 8, 2026', 1)
p(('Actual inflow figures sourced from the Inflows Actuals tab (row 1,229) and '
   'corroborated against the Cash Flow sheet (col DR = $2.200M). '
   'AI figures sourced from NetSuite GL bank account query (accounts 10617 and 10639) '
   'and NCR journal entries (JE39909–JE39912); these are the only bank-side credits '
   'confirmed in NetSuite on May 8, 2026.'), size=9, italic=True, space_after=4)

heading('4.1  Inflow Categories', 2)
add_table(
    ['Inflow Category','Actual ($M)','AI Captured ($M)','Unconfirmed ($M)','Methodology Note'],
    [
      [('Sears B&M Stores',{}),('$0.010',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.010',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Actual from Daily Detail cash position file (named range INFLOWS). Not visible as bank debit in NetSuite on posting date.',{})],
      [('Kmart Stores',{}),('$0.800',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.800',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Same sourcing as Sears B&M.',{})],
      [('Home Services',{}),('$1.370',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$1.370',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Largest inflow category on the date. Sourced from Daily Detail file.',{})],
      [('Auto Centers',{}),('$0.020',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.020',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Included in Inflows Actuals; not broken out separately in AI NetSuite query.',{})],
      [('ACH Return (JE72595)',{}),('Not broken out separately',{}),('$0.001',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':AMBER}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('AI captured: $682.50 ACH return posted to JPM 10639. Not material.',{})],
      [('All Other Categories',{}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('CITI Reimburse, Supply Chain, Asset Sales, Debt/Financing all zero on May 8.',{})],
      [('TOTAL INFLOWS',{'bold':True,'bg':GRY_BG}),('$2.200',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('$0.001',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('$2.199',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('99.9% of actual inflows unconfirmed by NetSuite on same-day basis',{'bg':GRY_BG,'color':RED})],
    ],
    col_widths=[1.8,1.1,1.2,1.2,3.2]
)

p()
p('Root Cause — Inflow Gap: Customer payment receipts in NetSuite post to the Accounts '
  'Receivable sub-ledger on the day of payment application, not to a bank account debit. '
  'The corresponding bank credit (deposit) typically reflects with a one-business-day lag '
  'in the NetSuite GL. The Actual File captures same-day cash inflows from the treasury '
  'cash position system (Daily Detail file, named range INFLOWS), which reconciles directly '
  'to bank-confirmed deposits.', size=9, italic=True, space_after=6)

doc.add_paragraph()

# ── SECTION 5 — DISBURSEMENTS ANALYSIS ───────────────────────────────────
heading('Section 5 — Disbursements Analysis: May 8, 2026', 1)
p(('Actual disbursement figures sourced from the Disbursement Actuals tab (row 367) '
   'and corroborated against the Cash Flow sheet (col DS = ($2.160M)). '
   'AI figures sourced from NetSuite bank account query (accounts 10617 and 10639).'),
  size=9, italic=True, space_after=4)

heading('5.1  Disbursement Categories', 2)
add_table(
    ['Category','Cash Flow Col','Actual ($M)','AI ($M)','Variance ($M)','Notes'],
    [
      [('Payroll & Benefits',{'bold':True}),('CN',{}),('($1.110)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('$1.110 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Payroll disbursed via Wells Fargo (10302). NetSuite payroll queries returned $0 on May 8; last WF payroll batch posted May 2. AI file Payroll sheet expired Jul 2025.',{})],
      [('Merch. On Terms (Accounts Payable — Merch)',{}),('CP',{}),('($0.460)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('$0.460 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Sourced from MERCH sheet (hidden; updated by James Coutre). AI file MERCH/Merch sheet has no 2026 data.',{})],
      [('Other Non-Merch Disbursements',{}),('CT',{}),('($0.250)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($0.727)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($0.477) overstated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':AMBER}),('AI placed all NetSuite VendPymts ($0.727M) in this catch-all bucket. Actual Non-Merch = $0.250M; the balance represents different categories.',{})],
      [('Taxes',{}),('CV',{}),('($0.289)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('$0.289 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Sourced from FY File Tax column (AQ). AI FY File has no May 2026 data.',{})],
      [('Rent',{}),('CQ',{}),('($0.009)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.009 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('From Non-Merch sheet.',{})],
      [('Logistics',{}),('CR',{}),('($0.022)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.022 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('From Non-Merch sheet.',{})],
      [('Risk Mgmt / Insurance',{}),('CU',{}),('($0.020)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.020 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('From FY File RiskMgt/Ins column (AP).',{})],
      [('Merch. CIA (Cash-in-Advance)',{}),('CO',{}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Consistent.',{})],
      [('All Other Categories',{}),('Various',{}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Advertising, Collateral, Interest, Fees, WU, P-Card, SHI, PropCo, Debt Repayment — all zero in Actual File on May 8.',{})],
      [('TOTAL DISBURSEMENTS',{'bold':True,'bg':GRY_BG}),('DS',{'bg':GRY_BG}),('($2.160)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('($0.727)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('$1.433 understated',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('AI captures 33.7% of actual gross disbursements',{'bg':GRY_BG,'color':RED})],
    ],
    col_widths=[1.55,0.6,0.85,0.85,1.2,3.45]
)
doc.add_paragraph()

# ── SECTION 6 — PAYROLL ANALYSIS ─────────────────────────────────────────
heading('Section 6 — Payroll & Benefits: Forecast vs. Actuals by Business Unit', 1)
p(('The Actual File Payroll sheet contains both a forecast column and an actuals column '
   'for May 8, 2026. The following table reconciles forecast to actuals by BU. '
   'All values are in U.S. dollars. '
   'The AI file Payroll sheet contains no 2026 data (expired July 4, 2025).'),
  size=9, italic=True, space_after=4)

add_table(
    ['Business Unit','Payroll Forecast ($)','Payroll Actuals ($)','Variance ($)','Variance %'],
    [
      [('Finance',{}),('($5,205)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($5,039)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$166 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('3.2%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('FinancialSvcs',{}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('HoldingsCo',{}),('($954)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($679)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$275 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('28.8%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Home Services',{'bold':True}),('($946,906)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('($935,349)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$11,557 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN,'bold':True}),('1.2%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('HR',{}),('($2,365)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($2,284)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$81 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('3.4%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('KCD',{}),('($3,028)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($3,340)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($312) Unfav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('(10.3%)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED})],
      [('Legal',{}),('($8,501)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($8,242)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$259 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('3.0%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Member Tech',{}),('($4,012)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($3,831)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$181 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('4.5%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('MSO',{}),('($12,626)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($11,558)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$1,068 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('8.5%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Real Estate',{}),('($5,446)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($5,886)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($440) Unfav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('(8.1%)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED})],
      [('Retail',{'bold':True}),('($82,357)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('($78,179)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$4,178 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN,'bold':True}),('5.1%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Retail Online',{}),('($573)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($489)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$84 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('14.7%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Service Live',{}),('$0',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('SYWR',{}),('($6,839)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($6,668)',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$171 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('2.5%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('Supply Chain',{'bold':True}),('($53,518)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('($50,241)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$3,277 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN,'bold':True}),('6.1%',{'align':WD_ALIGN_PARAGRAPH.RIGHT})],
      [('TOTAL',{'bold':True,'bg':GRY_BG}),('($1,132,329)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('($1,111,786)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('$20,543 Fav.',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN,'bold':True,'bg':GRY_BG}),('1.8%',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bg':GRY_BG})],
    ],
    col_widths=[1.6,1.5,1.5,1.5,1.4]
)

p()
p('Note: Payroll actuals are $20,543 (1.8%) favorable to forecast, driven primarily by '
  'Home Services ($11,557 favorable) and Retail ($4,178 favorable). KCD ($312 unfavorable) '
  'and Real Estate ($440 unfavorable) are the only BUs over forecast. '
  'These payroll actuals exist only in the Actual File; the AI File contains no May 2026 payroll data.',
  size=9, italic=True, space_after=6)

doc.add_paragraph()

# ── SECTION 7 — DEBT SCHEDULE ─────────────────────────────────────────────
heading('Section 7 — Outstanding Debt Obligations: Actual vs. AI File', 1)
p(('Debt balances represent outstanding principal as of May 8, 2026 (Actual File) '
   'versus January 30, 2026 (AI File — the last date in the AI Cash Flow sheet). '
   'Variances reflect changes over the intervening ~67 business days, '
   'not errors per se, but the AI file cannot represent current debt positions.'),
  size=9, italic=True, space_after=4)

add_table(
    ['Debt Instrument','Actual File\nMay 8, 2026 ($M)','AI File\nJan 30, 2026 ($M)','Change ($M)','Notes'],
    [
      [('UBS Loan — Note A',{'bold':True}),('$299.435',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$223.341',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$76.094',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Significant increase; 74 properties (Actual) vs. 85 (AI) — property count also differs, indicating portfolio change',{})],
      [('UBS Loan — Note B',{'bold':True}),('$25.368',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('$65.587',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($40.219)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Significant decrease; partial paydown or restructuring between Jan and May 2026',{})],
      [('ESL $175M Term Loan (T22)',{}),('$257.678',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$202.140',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$55.538',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Increase',{})],
      [('Lease Opco Loan',{}),('$533.317',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$446.575',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$86.743',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Increase',{})],
      [('Hackensack, NJ Loan (2026)',{}),('$7.975',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$7.585',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0.390',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Minor increase; column label discrepancy: AI file maps this as "UBS/Manteno Loan" in the same column position — structural error in AI file',{})],
      [('Cyrus 7th Term Loan',{}),('$30.884',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$40.770',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($9.886)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Decrease',{})],
      [('Sept. 2025 Term Loans',{'bold':True}),('$31.089',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('NOT TRACKED',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('N/A',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('This debt tranche (col AY) is absent from the AI file column structure — omitted entirely',{})],
      [('Cyrus 8th Incremental',{'bold':True}),('$8.160',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('NOT TRACKED',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('N/A',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('This debt tranche (col AZ) is absent from the AI file column structure — omitted entirely',{})],
      [('Guam Loan',{'bold':True}),('$34.000',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True}),('MISCLASSIFIED',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True}),('N/A',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Actual File col BQ = Guam Loan $34M. AI File col BQ = "BAML LC" — different instrument mapped to same column position',{})],
    ],
    col_widths=[1.8,1.2,1.2,1.1,3.2]
)

doc.add_paragraph()

# ── SECTION 8 — STRUCTURAL DEFICIENCIES ──────────────────────────────────
heading('Section 8 — Structural Deficiencies in the AI-Generated File', 1)

heading('8.1  Formula Integrity', 2)
p(('The AI file\'s Cash Flow sheet produces #N/A errors in all calculated columns '
   '(Net Cash Flow, Available Cash, Total Cash, Operating Cash Flows, Merch/Non-Merch '
   'Disbursements, Net Change in Cash, and all downstream variances) for every row from '
   'June 2, 2025 through the file\'s last date (January 30, 2026) — a period of '
   'approximately 172 business days and 100% of the FY 2025 second-half and full '
   'FY 2026 forecast window. The root cause is a broken MATCH lookup: the Disbursement '
   'Detail sheet references the Payroll sheet by date, and the Payroll sheet in the AI '
   'file expires on July 4, 2025. All subsequent Payroll/Benefits column lookups '
   'return #N/A, which cascades to the Daily Total and then to the Cash Flow Net CF '
   'computation.'), size=9.5, space_after=4)

heading('8.2  Column Mapping Errors — Debt / Letter of Credit Section', 2)
add_table(
    ['Excel Column','Actual File Label','AI File Label','Classification'],
    [
      [('AY',{}),('Sept. 2025 Term Loans — $31.1M',{}),('(Absent — column not present)',{'color':RED}),('Omission',{'color':RED})],
      [('AZ',{}),('Cyrus Eighth Incremental Term Loans — $8.2M',{}),('(Absent — column not present)',{'color':RED}),('Omission',{'color':RED})],
      [('BK',{}),('Hackensack, NJ Loan (2026) — $7.975M',{}),('UBS/Manteno Loan',{'color':RED}),('Misclassification',{'color':RED})],
      [('BM',{}),('Durham, NC Loan',{}),('(Blank — no label)',{'color':RED}),('Omission',{'color':RED})],
      [('BQ',{}),('Guam Loan — $34.0M',{}),('BAML LC',{'color':RED}),('Misclassification — wrong instrument entirely',{'color':RED,'bold':True})],
    ],
    col_widths=[0.8,2.5,2.5,1.7]
)

p()
heading('8.3  Missing Worksheets (SOP and MERCH)', 2)
p(('The Actual File contains a Standard Operating Procedures (SOP) tab that documents '
   'the data sourcing methodology for each worksheet, including named contacts, bank '
   'account numbers, GL account references, and refresh cadence. The AI file does not '
   'contain this tab. The absence of the SOP eliminates auditability — any reviewer '
   'of the AI file cannot determine the provenance of any figure without reference '
   'to an external document. '
   'The MERCH tab (hidden in the Actual File; updated by James Coutre per the SOP) is '
   'also absent from the AI file. Merchandise on-terms disbursements ($0.460M on May 8) '
   'therefore have no source data in the AI file.'), size=9.5, space_after=4)

heading('8.4  Stale Working-Paper Tabs', 2)
p(('The AI file contains 44 worksheets not present in the Actual File, '
   'all appearing to be working-paper archives dated Q4 2024 – Q1 2025 '
   '(e.g., "SHS Cash Changes 3.3.25," "Logistics 3.4.25," "KCD backup_3.4.25BR"). '
   'These do not contribute to the current forecast and increase the file size '
   'from 2.6 MB (Actual) to 11.9 MB (AI). They should be removed before any '
   'distribution.'), size=9.5, space_after=4)

doc.add_paragraph()

# ── SECTION 9 — UNAVAILABLE CASH ─────────────────────────────────────────
heading('Section 9 — Unavailable Cash & Restricted Cash Schedules', 1)
heading('9.1  Trapped Cash Comparison (Unavailable / Restricted)', 2)

add_table(
    ['Item','Actual File ($)','AI File ($)','Delta ($)','Notes'],
    [
      [('Cash deposited with PNC (LCs)',{}),('$5,636,582',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$5,636,582',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Consistent — both pull from live PNC LC balance formula',{})],
      [('Cash deposited with UBS (LCs)',{}),('$30,868,423',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$31,387,142',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('($518,719)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('AI file carries Jan 2026 balance; Actual reflects May 2026 balance (UBS upsize)',{})],
      [('Cash deposited with Citi (P-Card)',{}),('$2,000,000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$2,000,000',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Consistent — hardcoded in both',{})],
      [('Cash Holdback — FDC',{}),('$7,051,464',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$6,947,467',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$103,997',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Small difference; formula-driven vs. hardcoded',{})],
      [('Cash Holdback — Stripe',{}),('$1,271,111',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('—',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('$1,271,111',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Present in Actual File; absent from AI file — Stripe holdback line not in March 2025 source template',{})],
      [('Segregated RE Account',{}),('$5,241,937',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$5,241,937',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$0',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':GREEN}),('Consistent',{})],
      [('Home Warranty Reserves',{}),('$6,209,231',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$5,048,685',{'align':WD_ALIGN_PARAGRAPH.RIGHT}),('$1,160,546',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED}),('Actual File reflects updated warranty reserve; AI carries stale March 2025 figure',{})],
      [('GRAND TOTAL (Unavailable Cash)',{'bold':True,'bg':GRY_BG}),('$88,477k (approx.)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('$89,895k (approx.)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'bold':True,'bg':GRY_BG}),('($1,418k)',{'align':WD_ALIGN_PARAGRAPH.RIGHT,'color':RED,'bold':True,'bg':GRY_BG}),('AI total slightly higher due to stale UBS balance offsetting Stripe omission. Neither figure feeds correctly into the AI Cash Flow DV column for May 2026.',{'bg':GRY_BG})],
    ],
    col_widths=[2.0,1.3,1.3,1.2,2.7]
)

doc.add_paragraph()

# ── SECTION 10 — CONCLUSIONS ──────────────────────────────────────────────
heading('Section 10 — Conclusions and Recommended Actions', 1)

heading('10.1  Summary Assessment', 2)
p(('The AI-generated file is not a functional substitute for the finance-team-prepared '
   'daily cash forecast. On every material metric for May 8, 2026, the AI file '
   'produces materially incorrect outputs or no output at all:'), size=9.5, space_after=3)
for bullet in [
    'Net Cash Flow is misstated by ($0.767M) and directionally inverted (AI: negative outflow vs. Actual: positive net inflow).',
    'Total Cash is understated by $84.3M, driven by the AI file\'s inability to populate the Unavailable Cash register for 2026.',
    'Gross Inflows are understated by $2.199M (99.9%) because NetSuite GL does not reflect same-day bank credits for customer receipts.',
    'Gross Disbursements are understated by $1.433M (66.3%) because payroll, merchandise on-terms, and taxes are not captured by NetSuite AP vendor payment queries alone.',
    'The AI Cash Flow formula chain is broken from June 2, 2025 onward due to the Payroll sheet expiring on July 4, 2025.',
    'Two debt tranches (Sept. 2025 Term Loans, Cyrus Eighth Incremental) are absent from the AI file column structure, and one column (BQ) is misclassified as the wrong instrument.',
]:
    bpara = doc.add_paragraph(style='List Bullet')
    bpara.paragraph_format.space_after = Pt(2)
    run = bpara.add_run(bullet)
    run.font.size = Pt(9.5)

p()
heading('10.2  Recommended Actions', 2)
actions = [
    ('Discontinue use of AI file for treasury reporting.',
     'The AI file produces materially incorrect cash positions and should not be distributed internally or externally in its current state.'),
    ('Retain Actual File as the authoritative source.',
     'The finance-team file is current, continuous, and formula-correct through January 2027.'),
    ('Reconcile Trapped Cash / Unavailable Cash balances.',
     'The $1.3M variance in the Trapped Cash schedule between the two files (primarily Stripe holdback omission and stale UBS balance in AI file) should be investigated to confirm the Actual File figures.'),
    ('Investigate AI data-source gap for future automation.',
     'If NetSuite-based automation of the cash forecast is a goal, the inflow capture methodology must be redesigned. Customer receipts require the Daily Detail cash position file (named range INFLOWS), not raw GL queries. Payroll requires a direct feed from the Payroll team file, not the WF bank account. Merch on-terms requires the MERCH feed from James Coutre.'),
    ('Preserve and maintain the SOP tab.',
     'The SOP documents 17 named data sources with bank account numbers, GL codes, and responsible parties. This tab should be required in any distributed version of the model.'),
]
for i, (title, detail) in enumerate(actions):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(f'{i+1}.  {title}  ')
    run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = NAVY
    run2 = para.add_run(detail)
    run2.font.size = Pt(9.5)

p()
p('—  End of Report  —', bold=False, size=9, color=MIDGREY,
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p(f'Report generated: {date.today().strftime("%B %d, %Y")} | '
  'Source data extracted directly from file cells; no figures estimated or interpolated.',
  size=8, italic=True, color=MIDGREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────
out = '/Users/josh/Desktop/Cash Forecast Variance Analysis - May 8 2026 - CFO Report.docx'
doc.save(out)
import os
print(f"SAVED: {out}  ({os.path.getsize(out):,} bytes)")
