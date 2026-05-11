"""
Convert MASTER_INSIGHTS_DOCUMENT.md to a formatted Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, copy

MD_PATH = "/Users/josh/Downloads/SP_Analysis/MASTER_INSIGHTS_DOCUMENT.md"
OUT_PATH = "/Users/josh/Downloads/Transform SR - Master Business Insights Document.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level):
    """Add a heading with custom styling."""
    clean = text.strip().lstrip('#').strip()
    h = doc.add_heading(clean, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(clean)
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x6E, 0xA3)
        run.bold = True
    return h

def shade_cell(cell, hex_color):
    """Fill table cell background."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_inline_text(para, text):
    """Add text with **bold** and *italic* inline formatting, strip leading '- '."""
    # Remove leading '- ' or '* ' bullet markers
    text = re.sub(r'^[-\*]\s+', '', text)

    # Split on ** for bold and * for italic
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            r.italic = True
        else:
            para.add_run(part)

def build_table(doc, lines):
    """Parse a markdown table block and add a formatted Word table."""
    rows = []
    for ln in lines:
        if re.match(r'\|\s*[-:]+', ln):
            continue  # skip separator row
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.cell(ri, ci)
            # Strip bold markers for table cells — apply bold formatting instead
            is_bold = '**' in cell_text
            clean = re.sub(r'\*\*', '', cell_text).strip()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean)
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
                shade_cell(cell, '1F3A5F')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif is_bold:
                run.bold = True
                shade_cell(cell, 'E8F0F8')
            elif ri % 2 == 0:
                shade_cell(cell, 'F5F8FC')

    # Set column widths roughly equal
    total_width = Inches(6.5)
    col_w = int(total_width.emu / ncols)
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_w

    doc.add_paragraph()  # spacing after table


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    with open(MD_PATH, 'r') as f:
        raw = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    lines = raw.split('\n')
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # ── Heading 1 ──
        if line.startswith('# ') and not line.startswith('## '):
            set_heading(doc, line, 1)
            i += 1

        # ── Heading 2 ──
        elif line.startswith('## ') and not line.startswith('### '):
            set_heading(doc, line, 2)
            i += 1

        # ── Heading 3 ──
        elif line.startswith('### '):
            set_heading(doc, line, 3)
            i += 1

        # ── Horizontal rule ──
        elif line.strip() == '---':
            doc.add_paragraph()
            i += 1

        # ── Table block ──
        elif line.startswith('|'):
            table_lines = []
            while i < n and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            build_table(doc, table_lines)

        # ── Bullet / list item ──
        elif re.match(r'^[-\*]\s+', line) or re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_inline_text(p, line)
            i += 1

        # ── Blank line ──
        elif line.strip() == '':
            i += 1

        # ── Normal paragraph ──
        else:
            # Collect consecutive non-blank, non-special lines as one paragraph
            para_lines = []
            while i < n and lines[i].strip() != '' and not lines[i].startswith('#') \
                    and not lines[i].startswith('|') and not lines[i].strip() == '---' \
                    and not re.match(r'^[-\*]\s+', lines[i]) \
                    and not re.match(r'^\d+\.\s+', lines[i]):
                para_lines.append(lines[i])
                i += 1
            combined = ' '.join(para_lines).strip()
            if combined:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_inline_text(p, combined)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")

main()
